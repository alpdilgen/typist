"""
typist_core.py — Anova Typist Core Logic
=========================================
Used by both the Streamlit prototype and the FastAPI portal.
Dependencies: anthropic, python-docx, Pillow
"""

import base64
import html as _html
import io
import re
import uuid
from datetime import date
from typing import Optional

import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm
from PIL import Image

# ---------------------------------------------------------------------------
# Supported formats
# ---------------------------------------------------------------------------
SUPPORTED_FORMATS = {
    "pdf":  "application/pdf",
    "jpeg": "image/jpeg",
    "jpg":  "image/jpeg",
    "png":  "image/png",
    "webp": "image/webp",
    "tiff": "image/png",   # Converted to PNG via Pillow
    "tif":  "image/png",
    "bmp":  "image/png",   # Converted to PNG via Pillow
}

MAX_FILE_SIZE_MB = 20

# ---------------------------------------------------------------------------
# Typist Prompt (aligned with SKILL.md)
# ---------------------------------------------------------------------------
TYPIST_PROMPT_BASE = """You are a professional document digitization and transcription agent.
Convert the uploaded document into accurately formatted, machine-readable text while
preserving original layout, character formatting, and linguistic integrity.

## Step 1 — Pre-Transcription Analysis
Identify: page count, structure (headings/body/margins/columns), formatting
(bold/italic/underline/font sizes), layout (headers/footers/tables/lists),
document type, special characters, language(s), scan quality.

## Step 2 — Transcribe
Extract ALL visible text in reading order (left-to-right, top-to-bottom).

## Transcription Rules
- Extract ALL visible text in reading order
- Preserve spacing that indicates structure (line breaks, paragraph breaks)
- Use Markdown formatting: **bold**, *italic*, [UNDERLINED: text]
- Headings: # H1, ## H2, ### H3
- Tables: pipe format | Col1 | Col2 |
- Lists: - unordered, 1. ordered
- Page breaks: ---PAGE BREAK---
- Unclear word: [?word] | {image_rule}
- Uncertain handwriting: [HANDWRITTEN - UNCERTAIN: text]
- Illegible: [HANDWRITTEN - ILLEGIBLE]
- Language switch: [LANGUAGE SWITCH: Language]
- Forms: [FILLED: response] or [BLANK FIELD]
- NEVER correct spelling mistakes — preserve as-is
- Distinguish carefully: O vs 0, l vs 1, I vs 1
- Preserve all diacritics: é, ñ, ü, ş, ğ, č, etc.
- TWO-COLUMN TABLE LAYOUTS: Many forms and patent documents use a two-column table
  where field labels appear in the left column and values (or blank fields) in the right.
  ALWAYS detect and transcribe these as pipe-format Markdown tables, even if column
  borders are faint, implied, or obscured by watermarks. Do not skip any field.

## Step 3 — Output Format
Structure your response EXACTLY in these four sections with these exact headers:

### SECTION 1 — DOCUMENT METADATA
Document Type:         [Letter / Report / Form / Table / etc.]
Languages Identified:  [ISO code (Language name) — Confidence: High/Medium/Low]
Total Pages:           [X]
Scan Quality:          [Excellent / Good / Fair / Poor]
Handwritten Content:   [Yes (approx. X%) / No]
Transcription Date:    [YYYY-MM-DD]
Uncertain Elements:    [X flagged items]

### SECTION 2 — TRANSCRIBED CONTENT
[Full transcribed text with all formatting preserved]

### SECTION 3 — FORMATTING NOTES
[Formatting decisions, ambiguities resolved, layout choices]

### SECTION 4 — QUALITY NOTES
[Limitations, unclear sections, recommendations. If no issues: "Document transcribed with high confidence. No manual review required."]
"""

IMAGE_RULE_INCLUDE = "Image/diagram: [IMAGE: description]"
IMAGE_RULE_EXCLUDE = "Images/diagrams: DO NOT describe images or include any [IMAGE: ...] placeholders. Skip all visual elements silently."


def build_prompt(include_image_placeholders: bool = True) -> str:
    rule = IMAGE_RULE_INCLUDE if include_image_placeholders else IMAGE_RULE_EXCLUDE
    return TYPIST_PROMPT_BASE.format(image_rule=rule)


# ---------------------------------------------------------------------------
# Helper: Image conversion (TIFF, BMP → PNG)
# ---------------------------------------------------------------------------
def _convert_to_supported(file_bytes: bytes, ext: str) -> tuple[bytes, str]:
    """Converts TIFF/BMP files to PNG format supported by Claude."""
    if ext in ("tiff", "tif", "bmp"):
        img = Image.open(io.BytesIO(file_bytes))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue(), "image/png"
    return file_bytes, SUPPORTED_FORMATS[ext]


# ---------------------------------------------------------------------------
# Main Transcription Function
# ---------------------------------------------------------------------------
def transcribe_document(
    file_bytes: bytes,
    filename: str,
    api_key: str,
    model: str = "claude-sonnet-4-6",
    include_image_placeholders: bool = True,
) -> dict:
    """
    Transcribes a document using Claude Vision.

    Returns:
        {
            "raw":                      str,   # Claude's full response
            "metadata":                 str,   # Section 1 content
            "content":                  str,   # Section 2 — transcription
            "formatting_notes":         str,   # Section 3
            "quality_notes":            str,   # Section 4
            "filename":                 str,
            "model":                    str,
            "include_image_placeholders": bool,
        }
    Raises:
        ValueError: Unsupported format or file too large
        anthropic.APIError: API error
    """
    # --- Validation ---
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext not in SUPPORTED_FORMATS:
        raise ValueError(
            f"Unsupported format: .{ext}\n"
            f"Supported: {', '.join(f'.{e}' for e in SUPPORTED_FORMATS)}"
        )

    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise ValueError(f"File too large: {size_mb:.1f} MB (max {MAX_FILE_SIZE_MB} MB)")

    # --- Build content block ---
    client = anthropic.Anthropic(api_key=api_key)
    prompt = build_prompt(include_image_placeholders)

    if ext == "pdf":
        file_b64 = base64.standard_b64encode(file_bytes).decode()
        content_blocks = [
            {
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": file_b64,
                },
            },
            {"type": "text", "text": prompt},
        ]
    else:
        converted_bytes, media_type = _convert_to_supported(file_bytes, ext)
        file_b64 = base64.standard_b64encode(converted_bytes).decode()
        content_blocks = [
            {
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": media_type,
                    "data": file_b64,
                },
            },
            {"type": "text", "text": prompt},
        ]

    # --- Claude API call ---
    message = client.messages.create(
        model=model,
        max_tokens=8192,
        messages=[{"role": "user", "content": content_blocks}],
    )

    raw_text = message.content[0].text

    # --- Parse response ---
    sections = _parse_sections(raw_text, include_image_placeholders)
    sections["raw"] = raw_text
    sections["filename"] = filename
    sections["model"] = model
    sections["include_image_placeholders"] = include_image_placeholders

    return sections


# ---------------------------------------------------------------------------
# Response Parser
# ---------------------------------------------------------------------------
def _strip_code_fences(text: str) -> str:
    """
    Claude sometimes wraps the metadata section in ``` code blocks.
    This function removes all opening and closing ``` markers.
    Metadata never contains legitimate code blocks.
    """
    text = re.sub(r"```[a-zA-Z]*", "", text)
    return text.strip()


def _clean_html_entities(text: str) -> str:
    """Converts &nbsp; and other common HTML entities to clean characters."""
    replacements = {
        "&nbsp;":   " ",
        "&amp;":    "&",
        "&lt;":     "<",
        "&gt;":     ">",
        "&quot;":   '"',
        "&#39;":    "'",
        "&mdash;":  "—",
        "&ndash;":  "–",
        "&hellip;": "…",
    }
    for entity, char in replacements.items():
        text = text.replace(entity, char)
    return text


def _strip_image_tags(text: str) -> str:
    """Removes [IMAGE: ...] placeholders from transcription content."""
    return re.sub(r"\[IMAGE:[^\]]*\]", "", text).strip()


def _parse_sections(text: str, include_image_placeholders: bool = True) -> dict:
    """Extracts the 4 sections from Claude's response."""
    pattern = re.compile(
        r"###\s*SECTION\s*(\d)\s*[—–-]\s*[^\n]+\n(.*?)(?=###\s*SECTION\s*\d|$)",
        re.DOTALL | re.IGNORECASE,
    )
    found = {m.group(1): m.group(2).strip() for m in pattern.finditer(text)}

    # Clean metadata code fences
    metadata_raw = found.get("1", "Metadata could not be parsed.")
    metadata_clean = _strip_code_fences(metadata_raw)

    content = found.get("2", "Content could not be parsed.")
    if not include_image_placeholders:
        content = _strip_image_tags(content)

    return {
        "metadata":         _clean_html_entities(metadata_clean),
        "content":          _clean_html_entities(content),
        "formatting_notes": _clean_html_entities(found.get("3", "No formatting notes found.")),
        "quality_notes":    _clean_html_entities(found.get("4", "No quality notes found.")),
    }


def _extract_uncertain_count(metadata: str) -> int:
    """Parses the uncertain elements count from the metadata text."""
    match = re.search(r"Uncertain Elements\s*:\s*(\d+)", metadata, re.IGNORECASE)
    if match:
        return int(match.group(1))
    return 0


# ---------------------------------------------------------------------------
# Flagged Items Extractor
# ---------------------------------------------------------------------------
_FLAG_TYPE_PATTERN = re.compile(
    r"\[HANDWRITTEN\s*-\s*(UNCERTAIN|ILLEGIBLE)([^\]]*)\]"
    r"|\[(\?[^\]]*)\]"
    r"|\[BLANK FIELD\]"
    r"|\[LANGUAGE SWITCH:[^\]]*\]",
    re.IGNORECASE,
)

def _extract_flagged_items(content: str) -> list[dict]:
    """
    Scans transcription content and returns a list of flagged items.
    Each item: {"line_no": int, "flag_type": str, "content": str}
    """
    flagged = []
    for line_no, line in enumerate(content.splitlines(), start=1):
        matches = list(_FLAG_TYPE_PATTERN.finditer(line))
        if not matches:
            continue
        for m in matches:
            if m.group(1):  # HANDWRITTEN - UNCERTAIN/ILLEGIBLE
                flag_type = f"Handwritten – {m.group(1).capitalize()}"
            elif m.group(3):  # [?word]
                flag_type = "Uncertain reading"
            elif m.group(0).upper().startswith("[BLANK"):
                flag_type = "Blank field"
            else:
                flag_type = "Language switch"
            # Truncate long lines for readability in table
            display = line.strip()
            if len(display) > 120:
                display = display[:117] + "..."
            flagged.append({
                "line_no":   line_no,
                "flag_type": flag_type,
                "content":   display,
            })
    return flagged


# ---------------------------------------------------------------------------
# DOCX Generator
# ---------------------------------------------------------------------------
def create_docx(result: dict) -> bytes:
    """
    Generates a Word document from the transcription result.

    Document structure (designed for manual review against the original):
      1. Title + file info
      2. Document Information   — metadata table
      3. Formatting Notes       — layout decisions made during transcription
      4. Quality Notes          — overall quality assessment and limitations
      5. Flagged Items          — table of uncertain / illegible / blank elements
      6. Transcription          — full transcribed text

    Returns: DOCX file content as bytes.
    """
    doc = Document()

    # --- Page layout ---
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.5)

    # --- Anova Brand Palette ---
    ANOVA_CHARCOAL = RGBColor(0x3A, 0x3A, 0x3A)   # Primary text / headings
    ANOVA_CORAL    = RGBColor(0xE8, 0x5C, 0x4A)   # Accents, CTAs, highlights
    ANOVA_AMBER    = RGBColor(0xF7, 0x93, 0x1E)   # Warnings, uncertain elements
    ANOVA_TEAL     = RGBColor(0x4E, 0xCD, 0xC4)   # Secondary accents, dividers

    def add_heading(text: str, level: int = 1):
        p = doc.add_heading(text, level=level)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(4)
        for run in p.runs:
            run.font.color.rgb = ANOVA_CHARCOAL if level == 1 else ANOVA_CORAL
        return p

    def add_divider():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(0)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "4ECDC4")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_shaded_paragraph(text: str, fill_hex: str = "FFF3CD", text_color: RGBColor = None):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        pPr.append(shd)
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.bold = True
        if text_color:
            run.font.color.rgb = text_color
        return p

    content      = result.get("content", "")
    uncertain_count = _extract_uncertain_count(result.get("metadata", ""))
    flagged_items   = _extract_flagged_items(content)

    # =========================================================================
    # 1. TITLE
    # =========================================================================
    title = doc.add_heading("Anova Typist — Transcription Report", 0)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after  = Pt(2)
    for run in title.runs:
        run.font.color.rgb = ANOVA_CHARCOAL

    subtitle = doc.add_paragraph(f"File: {result.get('filename', 'unknown')}")
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after  = Pt(4)
    subtitle.runs[0].font.size = Pt(10)
    subtitle.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    subtitle.runs[0].font.italic = True

    add_divider()

    # =========================================================================
    # 2. DOCUMENT INFORMATION
    # =========================================================================
    add_heading("1. Document Information", level=1)

    meta_lines = [ln for ln in result.get("metadata", "").splitlines() if ":" in ln and ln.strip()]
    if meta_lines:
        tbl = doc.add_table(rows=len(meta_lines), cols=2)
        tbl.style = "Table Grid"
        for r_idx, line in enumerate(meta_lines):
            key, _, val = line.partition(":")
            key = key.strip()
            val = val.strip()
            cell_key = tbl.cell(r_idx, 0)
            cell_val = tbl.cell(r_idx, 1)
            cell_key.text = key
            cell_val.text = val
            for run in cell_key.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = ANOVA_CHARCOAL
            for run in cell_val.paragraphs[0].runs:
                run.font.size = Pt(10)
                if "Uncertain" in key and uncertain_count > 0:
                    run.font.color.rgb = ANOVA_AMBER
                    run.bold = True
    else:
        p = doc.add_paragraph(result.get("metadata", ""))
        if p.runs:
            p.runs[0].font.size = Pt(10)

    add_divider()

    # =========================================================================
    # 3. FORMATTING NOTES
    # =========================================================================
    add_heading("2. Formatting Notes", level=1)
    fn_text = result.get("formatting_notes", "")
    if fn_text:
        p = doc.add_paragraph(fn_text)
        p.paragraph_format.space_after = Pt(4)
        if p.runs:
            p.runs[0].font.size = Pt(10)

    add_divider()

    # =========================================================================
    # 4. QUALITY NOTES
    # =========================================================================
    add_heading("3. Quality Notes", level=1)
    qn_text = result.get("quality_notes", "")
    if qn_text:
        p = doc.add_paragraph(qn_text)
        p.paragraph_format.space_after = Pt(4)
        if p.runs:
            p.runs[0].font.size = Pt(10)

    add_divider()

    # =========================================================================
    # 5. FLAGGED ITEMS
    # =========================================================================
    add_heading("4. Flagged Items", level=1)

    if flagged_items:
        add_shaded_paragraph(
            f"⚠  {len(flagged_items)} item(s) require manual review. "
            f"Each flagged line is listed below and highlighted in amber in the transcription.",
            fill_hex="FFF3CD",
            text_color=ANOVA_AMBER,
        )

        # Table: # | Flag Type | Content
        hdr_labels = ["#", "Flag Type", "Content"]
        tbl = doc.add_table(rows=1 + len(flagged_items), cols=3)
        tbl.style = "Table Grid"

        # Header row
        hdr_cells = tbl.rows[0].cells
        for col_idx, label in enumerate(hdr_labels):
            hdr_cells[col_idx].text = label
            for run in hdr_cells[col_idx].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tc = hdr_cells[col_idx]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "3A3A3A")
            tcPr.append(shd)

        # Data rows
        for row_idx, item in enumerate(flagged_items, start=1):
            row_cells = tbl.rows[row_idx].cells
            row_cells[0].text = str(item["line_no"])
            row_cells[1].text = item["flag_type"]
            row_cells[2].text = item["content"]
            for col_idx in range(3):
                for run in row_cells[col_idx].paragraphs[0].runs:
                    run.font.size = Pt(9)
                    if col_idx == 1:
                        run.font.color.rgb = ANOVA_AMBER
                        run.bold = True
            if row_idx % 2 == 0:
                for col_idx in range(3):
                    tc = row_cells[col_idx]._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "F4F4F4")
                    tcPr.append(shd)
    else:
        p = doc.add_paragraph("No flagged items — transcription completed with full confidence.")
        if p.runs:
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(0x55, 0x99, 0x55)

    add_divider()

    # =========================================================================
    # 6. TRANSCRIPTION
    # =========================================================================
    add_heading("5. Transcription", level=1)

    if not result.get("include_image_placeholders", True):
        note_p = doc.add_paragraph()
        note_run = note_p.add_run("Note: Image descriptions have been omitted per user settings.")
        note_run.italic = True
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    _add_markdown_content(doc, content)

    # =========================================================================
    # FOOTER
    # =========================================================================
    footer_p = doc.add_paragraph(
        f"Anova Translation  |  Generated: {date.today().isoformat()}"
        f"  |  Model: {result.get('model', 'claude-sonnet-4-6')}"
    )
    footer_p.paragraph_format.space_before = Pt(8)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        run.font.italic = True

    # --- Return as bytes ---
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Markdown → DOCX helper
# ---------------------------------------------------------------------------
UNCERTAIN_PATTERN = re.compile(
    r"\[HANDWRITTEN\s*-\s*(UNCERTAIN|ILLEGIBLE)[^\]]*\]|\[\?[^\]]*\]|\[BLANK FIELD\]",
    re.IGNORECASE,
)

def _add_markdown_content(doc: Document, text: str):
    """Converts Markdown-formatted text into DOCX paragraphs."""
    ANOVA_AMBER = RGBColor(0xF7, 0x93, 0x1E)  # Brand amber — uncertain element highlights
    lines = text.splitlines()
    i = 0
    table_buffer = []

    while i < len(lines):
        line = lines[i]

        # Page break
        if "---PAGE BREAK---" in line or "[PAGE" in line.upper():
            doc.add_page_break()
            i += 1
            continue

        # Table start
        if line.strip().startswith("|"):
            table_buffer.append(line)
            i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_buffer.append(lines[i])
                i += 1
            _add_table(doc, table_buffer)
            table_buffer = []
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=3)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=4)

        # Unordered list
        elif line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _apply_inline_formatting(p, line.strip()[2:])

        # Ordered list
        elif re.match(r"^\d+\.\s", line.strip()):
            p = doc.add_paragraph(style="List Number")
            _apply_inline_formatting(p, re.sub(r"^\d+\.\s", "", line.strip()))

        # Empty line
        elif not line.strip():
            doc.add_paragraph()

        # Normal paragraph — check if it contains uncertain flags
        else:
            p = doc.add_paragraph()
            if UNCERTAIN_PATTERN.search(line):
                # Render entire line in warning color + italic
                run = p.add_run(line)
                run.font.size = Pt(10)
                run.font.color.rgb = ANOVA_AMBER
                run.italic = True
            else:
                _apply_inline_formatting(p, line)
                if p.runs:
                    p.runs[0].font.size = Pt(10)

        i += 1


def _apply_inline_formatting(paragraph, text: str):
    """Applies **bold** and *italic* inline Markdown formatting."""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _add_table(doc: Document, lines: list[str]):
    """Creates a DOCX table from Markdown table lines."""
    rows = [
        [cell.strip() for cell in line.strip().strip("|").split("|")]
        for line in lines
        if not re.match(r"^\|[-| :]+\|$", line.strip())
    ]
    if not rows:
        return

    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < col_count:
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_text
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True


# ---------------------------------------------------------------------------
# XML helpers for XLIFF generation
# ---------------------------------------------------------------------------
def _xml_escape(text: str) -> str:
    """Escapes special XML characters in element content."""
    return _html.escape(str(text), quote=False)

def _xml_attr(text: str) -> str:
    """Escapes special XML characters in attribute values."""
    return _html.escape(str(text), quote=True)


# ---------------------------------------------------------------------------
# Source language extractor
# ---------------------------------------------------------------------------
def _extract_source_language(metadata: str) -> str:
    """
    Parses the source language ISO code from the metadata 'Languages Identified' line.
    Returns 'und' (undetermined) if not found.
    """
    match = re.search(
        r"Languages?\s+Identified\s*:\s*([a-z]{2,3}(?:-[A-Z]{2})?)",
        metadata,
        re.IGNORECASE,
    )
    if match:
        return match.group(1)
    return "und"


# ---------------------------------------------------------------------------
# SRX-inspired sentence splitter
# Ref: https://www.gala-global.org/srx-10
# XLIFF spec defers segmentation to external standards (SRX / UAX #29).
# ---------------------------------------------------------------------------

# Abbreviations that should NOT trigger a sentence break
_ABBREV_RE = re.compile(
    r'\b(Mr|Mrs|Ms|Dr|Prof|Sr|Jr|St|vs|etc|e\.g|i\.e|Fig|Eq|No|Tab|Vol|Ref|'
    r'Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Oct|Nov|Dec|'
    r'Mon|Tue|Wed|Thu|Fri|Sat|Sun|Corp|Ltd|Inc|Co)\.',
    re.IGNORECASE,
)

# Break rule: sentence-ending .!? followed by whitespace + uppercase / quote / bracket
_BREAK_RE = re.compile(r'(?<=[.!?])\s+(?=[A-Z"\u201C\u2018\(\[])')


def _split_sentences(text: str) -> list:
    """
    SRX-inspired sentence splitter for CAT tool segmentation.

    Break rules  : period / exclamation / question + whitespace + uppercase
    Exception rules: common abbreviations, decimal numbers, initials
    Short text (<60 chars): returned as-is (single sentence likely).
    """
    text = text.strip()
    if not text:
        return []
    if len(text) < 60:
        return [text]

    _PH = "⊘"   # placeholder for masked periods (safe Unicode, never in normal text)

    # Mask abbreviation periods to prevent false breaks
    masked = _ABBREV_RE.sub(lambda m: m.group().replace(".", _PH), text)
    # Mask decimal numbers  (3.14, 1,500.00)
    masked = re.sub(r"(\d)\.(\d)", lambda m: m.group(1) + _PH + m.group(2), masked)
    # Mask initials:  A. B. Smith
    masked = re.sub(r"\b([A-Z])\.\s+(?=[A-Z])", lambda m: m.group(1) + _PH + " ", masked)
    # Mask ellipsis  (…  or ...)
    masked = masked.replace("...", _PH * 3)

    parts = _BREAK_RE.split(masked)
    return [p.replace(_PH, ".").strip() for p in parts if p.strip()]


# ---------------------------------------------------------------------------
# Block extractor + sentence-level segmenter
# Returns list of dicts: {block_idx, type, sentences}
# type: 'paragraph' | 'heading' | 'list' | 'table'
# IDs in XLIFF: b{block_idx}-s{sent_idx}  →  used for apply-back
# ---------------------------------------------------------------------------

def _blocks_for_xliff(content: str) -> list:
    """
    Parses markdown transcription into translatable blocks with sentence splitting.

    - Headings, list items, table cells → single segment each (already short)
    - Regular paragraphs              → split into sentences via SRX rules
    """
    blocks = []
    current_para_lines: list = []

    def flush_para():
        if current_para_lines:
            full_text = " ".join(current_para_lines).strip()
            if len(full_text) > 1:
                blocks.append({"type": "paragraph", "sentences": _split_sentences(full_text)})
            current_para_lines.clear()

    for line in content.splitlines():
        stripped = line.strip()

        if not stripped:
            flush_para()
            continue

        if "---PAGE BREAK---" in stripped or re.match(r"^\[PAGE", stripped, re.IGNORECASE):
            flush_para()
            continue

        if re.match(r"^\|[\s\-|:]+\|$", stripped):   # table separator
            continue

        if stripped.startswith("|"):                   # table row
            flush_para()
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            row_text = " | ".join(c for c in cells if c)
            if len(row_text) > 1:
                blocks.append({"type": "table", "sentences": [row_text]})
            continue

        if stripped.startswith("#"):                   # heading
            flush_para()
            heading_text = re.sub(r"^#+\s*", "", stripped)
            if len(heading_text) > 1:
                blocks.append({"type": "heading", "sentences": [heading_text]})
            continue

        if re.match(r"^(\d+\.|[-*•])\s+", stripped):  # list item
            flush_para()
            if len(stripped) > 1:
                blocks.append({"type": "list", "sentences": [stripped]})
            continue

        current_para_lines.append(stripped)

    flush_para()

    for idx, block in enumerate(blocks):
        block["block_idx"] = idx

    return blocks


# ---------------------------------------------------------------------------
# Bilingual editor helpers
# ---------------------------------------------------------------------------

def get_segments_for_editor(content: str) -> list:
    """
    Returns a flat list of segments for the bilingual editor, each including
    the page number derived from ---PAGE BREAK--- markers in the transcription.

    Each segment dict:
        id         : 'b{block_idx}-s{sent_idx}'
        page       : 1-based page number
        block_idx  : global block index
        sent_idx   : sentence index within block
        type       : 'paragraph' | 'heading' | 'list' | 'table'
        text       : segment text
    """
    blocks      = []
    current_page = 1
    para_lines: list = []
    para_page   = 1

    def flush_para():
        if para_lines:
            full_text = " ".join(para_lines).strip()
            if len(full_text) > 1:
                blocks.append({"type": "paragraph",
                                "sentences": _split_sentences(full_text),
                                "page": para_page})
            para_lines.clear()

    for line in content.splitlines():
        stripped = line.strip()

        if not stripped:
            flush_para()
            continue

        if "---PAGE BREAK---" in stripped or re.match(r"^\[PAGE", stripped, re.IGNORECASE):
            flush_para()
            current_page += 1
            continue

        if re.match(r"^\|[\s\-|:]+\|$", stripped):        # table separator
            continue

        if stripped.startswith("|"):                        # table row
            flush_para()
            cells    = [c.strip() for c in stripped.strip("|").split("|")]
            row_text = " | ".join(c for c in cells if c)
            if len(row_text) > 1:
                blocks.append({"type": "table", "sentences": [row_text], "page": current_page})
            continue

        if stripped.startswith("#"):                        # heading
            flush_para()
            heading_text = re.sub(r"^#+\s*", "", stripped)
            if len(heading_text) > 1:
                blocks.append({"type": "heading", "sentences": [heading_text], "page": current_page})
            continue

        if re.match(r"^(\d+\.|[-*•])\s+", stripped):       # list item
            flush_para()
            if len(stripped) > 1:
                blocks.append({"type": "list", "sentences": [stripped], "page": current_page})
            continue

        para_lines.append(stripped)
        para_page = current_page

    flush_para()

    segments = []
    for idx, block in enumerate(blocks):
        for si, sentence in enumerate(block["sentences"]):
            if sentence.strip():
                segments.append({
                    "id":        f"b{idx}-s{si}",
                    "page":      block["page"],
                    "block_idx": idx,
                    "sent_idx":  si,
                    "type":      block["type"],
                    "text":      sentence,
                })

    # Normalise page numbers so they always start at 1.
    # Claude sometimes emits ---PAGE BREAK--- before the first page's content
    # which would push all segments to pages 2+ — subtract the offset.
    if segments:
        min_page = min(s["page"] for s in segments)
        if min_page > 1:
            offset = min_page - 1
            for s in segments:
                s["page"] -= offset

    return segments


def reconstruct_content_from_segments(segments: list) -> str:
    """
    Reassembles edited segments back into markdown-formatted content,
    re-inserting ---PAGE BREAK--- markers where the page number changes.

    Used by the bilingual editor's Save & Export to regenerate DOCX/XLIFF
    from corrected segment texts.
    """
    blocks_data: dict = {}
    for seg in segments:
        bi = seg["block_idx"]
        if bi not in blocks_data:
            blocks_data[bi] = {
                "type":      seg["type"],
                "page":      seg["page"],
                "sentences": {},
            }
        blocks_data[bi]["sentences"][seg["sent_idx"]] = seg["text"]

    lines      = []
    prev_page  = 1

    for bi in sorted(blocks_data.keys()):
        block      = blocks_data[bi]
        block_type = block["type"]
        block_page = block["page"]
        sents      = [block["sentences"][si] for si in sorted(block["sentences"].keys())]
        full_text  = " ".join(s for s in sents if s.strip())

        if not full_text.strip():
            continue

        # Insert page break(s) when page changes
        if block_page > prev_page:
            for _ in range(block_page - prev_page):
                lines.append("---PAGE BREAK---")
            prev_page = block_page

        if block_type == "heading":
            lines.append(f"# {full_text}")
        elif block_type == "list":
            lines.append(full_text)
        elif block_type == "table":
            lines.append(f"| {full_text} |")
        else:
            lines.append(full_text)

        lines.append("")   # blank line between blocks

    return "\n".join(lines)


# Keep old helper for any internal callers
def _segment_for_xliff(content: str) -> list:
    """Flat list of segment strings — kept for backwards compatibility."""
    blocks = _blocks_for_xliff(content)
    return [s for b in blocks for s in b["sentences"] if len(s.strip()) > 1]


# ---------------------------------------------------------------------------
# XLIFF 1.2 — Universal bilingual format
# Spec: https://docs.oasis-open.org/xliff/v1.2/os/xliff-core.html
# Compatible with ALL major CAT tools: SDL Trados, memoQ, Phrase, Wordfast,
#   Déjà Vu, OmegaT, Memsource, MateCat, and any XLIFF-capable tool.
# ---------------------------------------------------------------------------
def create_xliff(result: dict, target_language: str, source_language: str = "",
                 docx_bytes: bytes = None) -> bytes:
    """
    Generates a standard XLIFF 1.2 bilingual file with sentence-level segmentation.

    Segmentation:
      - Headings, list items, table rows → single segment
      - Regular paragraphs → split into sentences (SRX-inspired rules)
      - Segment IDs: b{block_idx}-s{sent_idx}  →  used by apply_xliff_to_docx()

    Skeleton:
      - If docx_bytes provided, the Word document is embedded as base64 skeleton.
      - This enables CAT tools and apply_xliff_to_docx() to reconstruct the
        translated Word document after translation is complete.

    source_language: BCP-47 code (e.g. 'en-US'). Required.
    target_language: BCP-47 code (e.g. 'tr-TR'). Required.
    docx_bytes: Word document bytes to embed as skeleton (optional but recommended).
    Returns: XLIFF 1.2 file content as UTF-8 bytes.
    """
    import base64 as _b64

    filename = result.get("filename", "document")
    blocks   = _blocks_for_xliff(result.get("content", ""))
    today    = date.today().isoformat()
    model    = result.get("model", "claude-sonnet-4-6")

    lines = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<xliff version="1.2" xmlns="urn:oasis:names:tc:xliff:document:1.2">',
        f'  <file original="{_xml_attr(filename)}"',
        f'        source-language="{source_language}"',
        f'        target-language="{target_language}"',
        f'        datatype="plaintext">',
        '    <header>',
    ]

    # Embed Word document as skeleton for round-trip apply-back
    if docx_bytes:
        b64 = _b64.b64encode(docx_bytes).decode("ascii")
        lines += ['      <skl>', '        <internal-file form="base64">']
        for i in range(0, len(b64), 76):          # RFC 2045 line length
            lines.append("        " + b64[i:i + 76])
        lines += ['        </internal-file>', '      </skl>']

    lines += [
        '      <tool tool-id="anova-typist" tool-name="Anova Typist"',
        '            tool-version="1.0" tool-company="Anova Translation"/>',
        f'      <note>Transcribed by Anova Typist on {today} | Model: {model}</note>',
        '    </header>',
        '    <body>',
    ]

    for block in blocks:
        bi = block["block_idx"]
        for si, sentence in enumerate(block["sentences"]):
            if not sentence.strip():
                continue
            lines += [
                f'      <trans-unit id="b{bi}-s{si}" xml:space="preserve">',
                f'        <source>{_xml_escape(sentence)}</source>',
                f'      </trans-unit>',
            ]

    lines += [
        '    </body>',
        '  </file>',
        '</xliff>',
    ]
    return "\n".join(lines).encode("utf-8")


# ---------------------------------------------------------------------------
# Apply translated XLIFF back to the embedded Word skeleton
# ---------------------------------------------------------------------------

def apply_xliff_to_docx(xliff_bytes: bytes) -> bytes:
    """
    Applies translated segments from a completed XLIFF to produce a
    translated Word document.

    Approach: reconstruct the entire transcription from XLIFF segment texts
    (target where translated, source as fallback), then generate a fresh Word
    document via create_docx().  This is more robust than the previous
    paragraph-counter approach, which silently broke whenever the document
    contained tables (table cells are not in doc.paragraphs, causing every
    subsequent paragraph to be mapped to the wrong XLIFF block).

    Requirements:
      - XLIFF must contain <skl><internal-file form="base64"> skeleton
        (generated by create_xliff() with docx_bytes parameter)
      - At least one <target> element must be non-empty
      - Segment IDs must follow b{block_idx}-s{sent_idx} scheme

    Returns: translated Word document as bytes.
    """
    import base64 as _b64
    import xml.etree.ElementTree as ET

    NS       = "urn:oasis:names:tc:xliff:document:1.2"
    root     = ET.fromstring(xliff_bytes.decode("utf-8"))
    file_el  = root.find(f"{{{NS}}}file")
    hdr_el   = file_el.find(f"{{{NS}}}header")

    # ── Extract skeleton ──────────────────────────────────────────────────
    skl_el = hdr_el.find(f"{{{NS}}}skl")
    if skl_el is None:
        raise ValueError(
            "No skeleton found in XLIFF. "
            "Re-export from Anova Typist with the 'Export bilingual XLIFF' option enabled."
        )
    int_file_el = skl_el.find(f"{{{NS}}}internal-file")
    b64_content = "".join((int_file_el.text or "").split())
    skeleton_docx = _b64.b64decode(b64_content)

    # ── Collect segments: source text + target text ───────────────────────
    body_el        = file_el.find(f"{{{NS}}}body")
    segments       = []
    has_translation = False

    for tu in body_el.findall(f"{{{NS}}}trans-unit"):
        tu_id = tu.get("id", "")
        m = re.match(r"^b(\d+)-s(\d+)$", tu_id)
        if not m:
            continue
        bi, si = int(m.group(1)), int(m.group(2))

        src_el  = tu.find(f"{{{NS}}}source")
        tgt_el  = tu.find(f"{{{NS}}}target")
        source  = (src_el.text or "").strip() if src_el is not None else ""
        target  = (tgt_el.text or "").strip() if (tgt_el is not None and tgt_el.text) else ""

        if target:
            has_translation = True

        segments.append({
            "id":        tu_id,
            "block_idx": bi,
            "sent_idx":  si,
            "page":      1,
            "type":      "paragraph",   # type is inferred by reconstruct
            "text":      target if target else source,
        })

    if not has_translation:
        raise ValueError(
            "No translated segments found in XLIFF. "
            "Make sure the file has been fully translated before applying."
        )

    # ── Rebuild translated content ────────────────────────────────────────
    translated_content = reconstruct_content_from_segments(segments)

    # ── Extract metadata / notes from skeleton Word ───────────────────────
    from docx import Document as _Document
    orig_doc  = _Document(io.BytesIO(skeleton_docx))
    orig_meta = _extract_section_text(orig_doc, "1. Document Information")
    orig_fmt  = _extract_section_text(orig_doc, "2. Formatting Notes")
    orig_qual = _extract_section_text(orig_doc, "3. Quality Notes")

    src_lang  = file_el.get("source-language", "")
    tgt_lang  = file_el.get("target-language", "")
    filename  = file_el.get("original", "document")

    # Append translation info to metadata
    if src_lang or tgt_lang:
        orig_meta = (orig_meta + f"\nTranslation:           {src_lang} → {tgt_lang}").strip()

    result = {
        "content":                  translated_content,
        "metadata":                 orig_meta  or "Metadata not available.",
        "formatting_notes":         orig_fmt   or "See original document.",
        "quality_notes":            orig_qual  or "Translated document — review translation quality manually.",
        "filename":                 filename,
        "model":                    "claude-sonnet-4-6",
        "include_image_placeholders": True,
    }

    return create_docx(result)


def _extract_section_text(doc, heading_keyword: str) -> str:
    """
    Extracts plain-text content of the Word section whose heading contains
    heading_keyword, stopping at the next heading. Includes text from tables
    that fall within the section (using XML iteration for correct order).
    """
    from docx.text.paragraph import Paragraph as _Para
    from docx.table         import Table     as _Tbl

    lines      = []
    capturing  = False

    for child in doc.element.body:
        tag = child.tag.split("}")[1] if "}" in child.tag else child.tag

        if tag == "p":
            p = _Para(child, doc)
            is_hdg = p.style.name.startswith("Heading") or p.style.name.startswith("Title")
            if is_hdg and heading_keyword in p.text:
                capturing = True
                continue
            if is_hdg and capturing:
                break
            if capturing and p.text.strip():
                lines.append(p.text.strip())

        elif tag == "tbl" and capturing:
            tbl = _Tbl(child, doc)
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(": ".join(cells) if len(cells) > 1 else cells[0])

    # Deduplicate consecutive identical lines (merged table cells repeat)
    deduped = []
    for line in lines:
        if not deduped or line != deduped[-1]:
            deduped.append(line)
    return "\n".join(deduped)


